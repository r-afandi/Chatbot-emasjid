import io
from typing import List
from pypdf import PdfReader
from docx import Document
import os
import csv
from app.core.config import settings
import requests
from bs4 import BeautifulSoup
import xml.etree.ElementTree as ET

class FileProcessingService:
    def __init__(self):
        # Create upload directory if it doesn't exist
        if not os.path.exists(settings.FILE_STORAGE_PATH):
            os.makedirs(settings.FILE_STORAGE_PATH)
    
    def process_pdf(self, content: bytes) -> List[str]:
        """Extract text from PDF file and split into chunks"""
        reader = PdfReader(io.BytesIO(content))
        
        texts = []
        for page in reader.pages:
            txt = page.extract_text()
            if txt:
                texts.append(txt)
        
        # Split into chunks - improved chunking for tables
        chunks = []
        for t in texts:
            # Split by sentence endings and line breaks
            for chunk in t.replace('\n', '. ').split(". "):
                chunk = chunk.strip()
                if len(chunk) > 20:  # Minimum chunk size
                    chunks.append(chunk)
        
        return chunks
    
    def process_docx(self, content: bytes) -> List[str]:
        """Extract text from DOCX file including tables and split into chunks"""
        doc = Document(io.BytesIO(content))
        
        texts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    # Join cell contents with pipes for table structure
                    table_row = " | ".join(row_text)
                    texts.append(f"Table row: {table_row}")
        
        # Join and split into chunks
        full_text = " ".join(texts)
        chunks = []
        for chunk in full_text.split(". "):
            if len(chunk.strip()) > 20:  # Minimum chunk size
                chunks.append(chunk.strip())
        
        return chunks
    
    def process_csv(self, content: bytes) -> List[str]:
        """Extract text from CSV file and split into chunks"""
        text = content.decode('utf-8')
        reader = csv.reader(io.StringIO(text))
        
        chunks = []
        headers = None
        
        for i, row in enumerate(reader):
            if i == 0:
                headers = row  # First row as headers
                continue
            
            if headers and len(row) == len(headers):
                # Create structured text from CSV row
                row_text = ", ".join(f"{headers[j]}: {row[j]}" for j in range(len(headers)) if row[j].strip())
                if row_text:
                    chunks.append(f"Data entry: {row_text}")
            else:
                # Fallback for rows without headers
                row_text = ", ".join(cell.strip() for cell in row if cell.strip())
                if row_text:
                    chunks.append(f"CSV row: {row_text}")
        
        return chunks
    
    def save_file(self, content: bytes, filename: str) -> str:
        """Save file to storage and return file path"""
        file_path = os.path.join(settings.FILE_STORAGE_PATH, filename)
        with open(file_path, "wb") as f:
            f.write(content)
        return file_path

    def process_url(self, url: str) -> List[str]:
        """Extract text from a webpage and split into chunks"""
        try:
            # Tambahkan headers agar bot kita tidak diblokir oleh sistem keamanan website
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.text, 'html.parser')
            
            # Hapus elemen script, style, nav, footer, header agar AI fokus pada konten utama
            for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
                element.extract()
                
            # Ambil teks bersih
            text = soup.get_text(separator=' ', strip=True)
            
            # Pecah menjadi chunk / potongan kalimat
            chunks = []
            for chunk in text.replace('\n', '. ').split(". "):
                chunk = chunk.strip()
                if len(chunk) > 20:  # Abaikan kata-kata yang terlalu pendek
                    chunks.append(chunk)
                    
            return chunks
        except Exception as e:
            print(f"Error scraping URL {url}: {e}")
            return []

    def process_sitemap(self, sitemap_url: str) -> List[str]:
        """Extract all URLs from a sitemap XML"""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get(sitemap_url, headers=headers, timeout=15)
            response.raise_for_status()
            
            # Parse XML
            root = ET.fromstring(response.content)
            
            urls = []
            # Cari semua tag 'loc' (biasanya berisi URL) tanpa mempedulikan namespace
            for elem in root.iter():
                if 'loc' in elem.tag:
                    if elem.text and elem.text.startswith('http'):
                        urls.append(elem.text.strip())
                        
            return urls
        except Exception as e:
            print(f"Error parsing sitemap {sitemap_url}: {e}")
            return []

# Initialize the file processing service
file_processing_service = FileProcessingService()